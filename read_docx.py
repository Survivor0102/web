#!/usr/bin/env python3
"""
使用python-docx库读取docx文件内容，处理中文编码
"""
import sys
import io
from docx import Document

def read_docx_content(docx_path):
    """读取docx文件的所有段落和表格"""
    try:
        doc = Document(docx_path)

        # 使用StringIO捕获输出，便于编码处理
        output = io.StringIO()

        output.write("=" * 60 + "\n")
        output.write(f"文档: {docx_path}\n")
        output.write("=" * 60 + "\n")

        # 读取段落
        output.write("\n【段落内容】\n")
        for i, paragraph in enumerate(doc.paragraphs, 1):
            text = paragraph.text.strip()
            if text:
                output.write(f"{i:3d}. {text}\n")

        # 读取表格
        output.write("\n【表格内容】\n")
        for t, table in enumerate(doc.tables, 1):
            output.write(f"\n表格 {t}:\n")
            for i, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                if any(row_data):  # 只打印有内容的行
                    output.write(f"  行{i+1}: {' | '.join(row_data)}\n")

        # 提取所有文本
        all_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)

        full_text = '\n'.join(all_text)

        output.write("\n" + "=" * 60 + "\n")
        output.write("【完整文本摘要】\n")
        output.write("=" * 60 + "\n")
        output.write(full_text[:2000] + ("..." if len(full_text) > 2000 else ""))
        output.write(f"\n\n总字符数: {len(full_text)}")

        return output.getvalue(), full_text

    except Exception as e:
        return f"读取文档时出错: {e}", None

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("使用方法: python read_docx.py <docx文件路径>")
        sys.exit(1)

    docx_path = sys.argv[1]
    output_str, content = read_docx_content(docx_path)

    # 尝试多种编码输出
    try:
        # 直接打印（系统默认编码）
        print(output_str)
    except UnicodeEncodeError:
        # 如果失败，尝试GBK编码
        try:
            print(output_str.encode('gbk', errors='ignore').decode('gbk'))
        except:
            print("无法正确显示中文内容")